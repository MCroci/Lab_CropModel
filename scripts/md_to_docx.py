"""Convert ESAME_ESERCIZI.md to Word (.docx)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def strip_latex(text: str) -> str:
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\cdot", "·", text)
    text = re.sub(r"\\max", "max", text)
    text = re.sub(r"\\min", "min", text)
    text = re.sub(r"\\left|\\right", "", text)
    text = re.sub(r"\\,", " ", text)
    text = re.sub(r"\\!", "", text)
    text = re.sub(r"\\tfrac\{([^}]*)\}\{([^}]*)\}", r"(\1/\2)", text)
    text = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1/\2)", text)
    text = re.sub(r"\{,\}", ",", text)
    text = re.sub(r"\^\{([^}]*)\}", r"^\1", text)
    text = re.sub(r"_\{([^}]*)\}", r"_\1", text)
    text = re.sub(r"\\\(|\\\)|\\\[|\\\]", "", text)
    text = re.sub(r"\$", "", text)
    text = re.sub(r"\\approx", "≈", text)
    text = re.sub(r"\\propto", "∝", text)
    text = re.sub(r"\\exp", "exp", text)
    text = re.sub(r"\\quad", " ", text)
    text = re.sub(r"\\underline\{\\hspace\{[^}]*\}\}", "______", text)
    return text.strip()


def add_formatted_paragraph(doc: Document, line: str, style: str | None = None):
    line = strip_latex(line)
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part:
            p.add_run(part)
    return p


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    in_solutions = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(strip_latex(stripped[2:]), level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            title = strip_latex(stripped[3:])
            if title.lower().startswith("soluzioni"):
                in_solutions = True
            doc.add_heading(title, level=1)
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(strip_latex(stripped[4:]), level=2)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_separator(lines[i]):
                    table_lines.append(parse_table_row(lines[i]))
                i += 1
            if table_lines:
                cols = max(len(r) for r in table_lines)
                table = doc.add_table(rows=len(table_lines), cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(table_lines):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        table.rows[ri].cells[ci].text = strip_latex(cell_text)
                        if ri == 0:
                            for p in table.rows[ri].cells[ci].paragraphs:
                                for run in p.runs:
                                    run.bold = True
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            add_formatted_paragraph(doc, stripped[1:-1], style="Intense Quote")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            add_formatted_paragraph(doc, stripped, style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        if stripped.startswith("\\[") or stripped.startswith("\\("):
            math_lines = [stripped]
            i += 1
            while i < len(lines) and not (
                lines[i].strip().endswith("\\]") or lines[i].strip().endswith("\\)")
            ):
                math_lines.append(lines[i].strip())
                i += 1
            if i < len(lines):
                math_lines.append(lines[i].strip())
                i += 1
            p = doc.add_paragraph()
            run = p.add_run(strip_latex(" ".join(math_lines)))
            run.italic = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue

        add_formatted_paragraph(doc, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Creato: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "ESAME_ESERCIZI.md"
    out = root / "ESAME_ESERCIZI.docx"
    if len(sys.argv) > 1:
        md = Path(sys.argv[1])
    if len(sys.argv) > 2:
        out = Path(sys.argv[2])
    convert(md, out)
