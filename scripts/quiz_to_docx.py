"""Export quiz (open + closed) from data/quizData.ts to Word."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX

ROOT = Path(__file__).resolve().parents[1]
QUIZ_TS = ROOT / "data" / "quizData.ts"
LABELS = ["A", "B", "C", "D"]


def _unescape(s: str) -> str:
    return s.replace("\\'", "'").replace('\\"', '"')


def _parse_quoted_strings(s: str) -> list[str]:
    """Extract quoted strings ('...' or \"...\") from a fragment."""
    out: list[str] = []
    i = 0
    while i < len(s):
        if s[i] not in ("'", '"'):
            i += 1
            continue
        quote = s[i]
        i += 1
        buf: list[str] = []
        while i < len(s):
            if s[i] == "\\" and i + 1 < len(s):
                buf.append(s[i + 1])
                i += 2
                continue
            if s[i] == quote:
                out.append("".join(buf))
                i += 1
                break
            buf.append(s[i])
            i += 1
    return out


def parse_open_questions(content: str) -> list[dict]:
    start = content.find("export const OPEN_QUESTIONS")
    end = content.find("export const OPEN_ANSWERS")
    block = content[start:end]
    pattern = re.compile(
        r"\{\s*id:\s*'(D\d+)',\s*number:\s*(\d+),\s*section:\s*'((?:[^'\\]|\\.)*)',\s*text:\s*",
    )
    items = []
    for m in pattern.finditer(block):
        rest = block[m.end() :]
        strings = _parse_quoted_strings(rest[:500])
        if not strings:
            continue
        items.append(
            {
                "id": m.group(1),
                "number": int(m.group(2)),
                "section": _unescape(m.group(3)),
                "text": _unescape(strings[0]),
            }
        )
    return items


def parse_closed_questions(content: str) -> list[dict]:
    items = []
    for line in content.splitlines():
        if not line.strip().startswith("mc("):
            continue
        head = re.match(
            r"\s*mc\('(D\d+)',\s*(\d+),\s*'((?:[^'\\]|\\.)*)',\s*",
            line,
        )
        if not head:
            continue
        rest = line[head.end() :]
        all_str = _parse_quoted_strings(rest)
        if len(all_str) < 5:
            continue
        text = all_str[0]
        opts = all_str[1:5]
        tail = rest[rest.rfind("]") + 1 :]
        ci_m = re.search(r",\s*(\d+)\s*\)", tail)
        if not ci_m:
            continue
        items.append(
            {
                "id": head.group(1),
                "number": int(head.group(2)),
                "section": _unescape(head.group(3)),
                "text": _unescape(text),
                "options": [_unescape(o) for o in opts],
                "correctIndex": int(ci_m.group(1)),
            }
        )
    return items


def parse_open_answers(content: str) -> dict[str, str]:
    start = content.find("export const OPEN_ANSWERS")
    end = content.find("export const CLOSED_QUESTIONS")
    block = content[start:end]
    pattern = re.compile(r"(D\d+):\s*`([\s\S]*?)`(?:,|\s*\})", re.MULTILINE)
    return {m.group(1): m.group(2).strip() for m in pattern.finditer(block)}


def add_paragraph(doc: Document, text: str, *, bold=False, italic=False, style=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_rich_paragraph(doc: Document, text: str, style=None):
    """Simple **bold** parsing."""
    p = doc.add_paragraph(style=style)
    for i, part in enumerate(re.split(r"(\*\*[^*]+\*\*)", text)):
        if not part:
            continue
        run = p.add_run(part[2:-2] if part.startswith("**") else part)
        if part.startswith("**"):
            run.bold = True
    return p


def build_docx(out_path: Path) -> None:
    content = QUIZ_TS.read_text(encoding="utf-8")
    open_q = parse_open_questions(content)
    closed_q = parse_closed_questions(content)
    answers = parse_open_answers(content)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("Prova di Valutazione — Domande e Risposte", level=0)
    add_paragraph(
        doc,
        "Università Cattolica del Sacro Cuore · Modellistica Applicata alle Produzioni Vegetali\n"
        "Modulo Coltivazioni Erbacee — CropModel Lab",
        italic=True,
    )
    add_paragraph(
        doc,
        f"Domande aperte: {len(open_q)} · Domande chiuse: {len(closed_q)} · "
        "Fonte: data/quizData.ts (quiz nell'applicazione).",
        italic=True,
    )
    doc.add_paragraph()

    # --- PARTE A ---
    doc.add_heading("PARTE A — Domande a risposta aperta", level=1)
    current_section = None
    for q in open_q:
        if q["section"] != current_section:
            current_section = q["section"]
            doc.add_heading(current_section, level=2)
        doc.add_heading(f"{q['id']} — Domanda {q['number']}", level=3)
        add_rich_paragraph(doc, q["text"])
        ans = answers.get(q["id"], "")
        if ans:
            p = doc.add_paragraph()
            r = p.add_run("Risposta di riferimento. ")
            r.bold = True
            for line in ans.split("\n"):
                if line.strip():
                    add_rich_paragraph(doc, line.strip())
        doc.add_paragraph()

    # --- PARTE B ---
    doc.add_page_break()
    doc.add_heading("PARTE B — Domande a risposta chiusa", level=1)
    add_paragraph(
        doc,
        "Una sola opzione corretta per domanda. La risposta corretta è evidenziata.",
        italic=True,
    )
    current_section = None
    for q in closed_q:
        if q["section"] != current_section:
            current_section = q["section"]
            doc.add_heading(current_section, level=2)
        doc.add_heading(f"{q['id']} — Domanda {q['number']}", level=3)
        add_rich_paragraph(doc, q["text"])
        for i, opt in enumerate(q["options"]):
            label = LABELS[i]
            p = doc.add_paragraph(style="List Bullet")
            prefix = f"{label}) "
            run = p.add_run(prefix + opt)
            if i == q["correctIndex"]:
                run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                p.add_run("  ✓ corretta")
        doc.add_paragraph()

    doc.save(str(out_path))
    print(f"Creato: {out_path}")
    print(f"  Aperte: {len(open_q)}, Chiuse: {len(closed_q)}, Risposte aperte: {len(answers)}")


if __name__ == "__main__":
    out = ROOT / "QUIZ_DOMANDE_E_RISPOSTE.docx"
    if len(sys.argv) > 1:
        out = Path(sys.argv[1])
    build_docx(out)
